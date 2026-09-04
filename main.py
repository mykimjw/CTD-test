import io
import os
import re
import copy
import difflib
import docx
from docx.shared import RGBColor, Pt, Inches
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from typing import Optional
from fastapi import FastAPI, UploadFile, File, Form, Header, HTTPException
from fastapi.responses import StreamingResponse

app = FastAPI()

MY_API_KEY = os.getenv("MY_API_KEY", "my-secret-key-1234")

COLOR_RED = RGBColor(255, 0, 0)
COLOR_BLUE = RGBColor(0, 0, 255)
COLOR_BLACK = RGBColor(0, 0, 0)
HEX_HEADER_BG = "E6EEF8"

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def get_char_diff(old_text, new_text):
    matcher = difflib.SequenceMatcher(None, old_text, new_text)
    opcodes = matcher.get_opcodes()
    diff_runs = []
    for tag, i1, i2, j1, j2 in opcodes:
        if tag == 'equal':
            diff_runs.append(('equal', old_text[i1:i2]))
        elif tag == 'delete':
            diff_runs.append(('delete', old_text[i1:i2]))
        elif tag == 'insert':
            diff_runs.append(('insert', new_text[j1:j2]))
        elif tag == 'replace':
            diff_runs.append(('delete', old_text[i1:i2]))
            diff_runs.append(('insert', new_text[j1:j2]))

    merged_runs = []
    for tag, text in diff_runs:
        if not text:
            continue
        if merged_runs and merged_runs[-1][0] == tag:
            merged_runs[-1] = (tag, merged_runs[-1][1] + text)
        else:
            merged_runs.append((tag, text))
    return merged_runs

def is_heading(paragraph):
    text = paragraph.text.strip()
    if not text:
        return False
    heading_patterns = [
        r'^제\s*\d+\s*[조장항장회]\b', r'^(?:\d+\.)+\d*\s*', r'^[가-힣]\.\s*',
        r'^(?:\d+|[가-힣]|[a-zA-Z])\)\s*', r'^\(\s*(?:\d+|[가-힣]|[a-zA-Z])\s*\)\s*',
        r'^[①-⑳]', r'^\[[^\]]+\]', r'^(?:■|●|▲|◆|○|▶|◈|※|•|\*)\s*'
    ]
    return any(re.match(pattern, text) for pattern in heading_patterns)

def get_table_summary(table):
    rows_text = []
    for row in table.rows:
        row_text = " | ".join(cell.text.strip() for cell in row.cells)
        rows_text.append(row_text)
    return "\n".join(rows_text)

def iter_block_items(parent):
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, docx.table._Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Invalid parent type")

    for child in parent_elm.iterchildren():
        if isinstance(child, docx.oxml.text.paragraph.CT_P):
            yield docx.text.paragraph.Paragraph(child, parent)
        elif isinstance(child, docx.oxml.table.CT_Tbl):
            yield docx.table.Table(child, parent)

def apply_old_inline_diff(paragraph, diff_runs):
    p_elm = paragraph._p
    for run in list(paragraph.runs):
        p_elm.remove(run._r)
    for tag, text in diff_runs:
        if tag == 'equal':
            paragraph.add_run(text)
        elif tag == 'delete':
            run = paragraph.add_run(text)
            run.font.color.rgb = COLOR_RED

def apply_new_inline_diff(paragraph, diff_runs):
    p_elm = paragraph._p
    for run in list(paragraph.runs):
        p_elm.remove(run._r)
    for tag, text in diff_runs:
        if tag == 'equal':
            paragraph.add_run(text)
        elif tag == 'insert':
            run = paragraph.add_run(text)
            run.font.color.rgb = COLOR_BLUE

def color_entire_block(block, rgb_color):
    if isinstance(block, docx.text.paragraph.Paragraph):
        for run in block.runs:
            run.font.color.rgb = rgb_color
    else:
        for row in block.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = rgb_color

def match_replaced_ranges(blocks1_range, blocks2_range):
    paired = []
    min_len = min(len(blocks1_range), len(blocks2_range))
    for k in range(min_len):
        paired.append((blocks1_range[k], blocks2_range[k]))
    if len(blocks1_range) > min_len:
        for k in range(min_len, len(blocks1_range)):
            paired.append((blocks1_range[k], None))
    elif len(blocks2_range) > min_len:
        for k in range(min_len, len(blocks2_range)):
            paired.append((None, blocks2_range[k]))
    return paired

def compare_and_mark_tables(t1, t2):
    rows1_repr = [" | ".join(c.text.strip() for c in r.cells) for r in t1.rows]
    rows2_repr = [" | ".join(c.text.strip() for c in r.cells) for r in t2.rows]

    row_matcher = difflib.SequenceMatcher(None, rows1_repr, rows2_repr)
    row_opcodes = row_matcher.get_opcodes()

    for r_tag, ri1, ri2, rj1, rj2 in row_opcodes:
        if r_tag == 'equal':
            continue
        elif r_tag == 'delete':
            for idx in range(ri1, ri2):
                for cell in t1.rows[idx].cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = COLOR_RED
        elif r_tag == 'insert':
            for idx in range(rj1, rj2):
                for cell in t2.rows[idx].cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.color.rgb = COLOR_BLUE
        elif r_tag == 'replace':
            old_rows, new_rows = t1.rows[ri1:ri2], t2.rows[rj1:rj2]
            min_rows = min(len(old_rows), len(new_rows))
            for k in range(min_rows):
                r1, r2 = old_rows[k], new_rows[k]
                min_cells = min(len(r1.cells), len(r2.cells))
                for c_idx in range(min_cells):
                    cell1, cell2 = r1.cells[c_idx], r2.cells[c_idx]
                    p1_list, p2_list = cell1.paragraphs, cell2.paragraphs
                    min_p = min(len(p1_list), len(p2_list))
                    for p_idx in range(min_p):
                        p1, p2 = p1_list[p_idx], p2_list[p_idx]
                        diff_runs = get_char_diff(p1.text.strip(), p2.text.strip())
                        apply_old_inline_diff(p1, diff_runs)
                        apply_new_inline_diff(p2, diff_runs)

                    if len(p1_list) > min_p:
                        for p_idx in range(min_p, len(p1_list)):
                            for r in p1_list[p_idx].runs: r.font.color.rgb = COLOR_RED
                    if len(p2_list) > min_p:
                        for p_idx in range(min_p, len(p2_list)):
                            for r in p2_list[p_idx].runs: r.font.color.rgb = COLOR_BLUE

            if len(old_rows) > min_rows:
                for idx in range(min_rows, len(old_rows)):
                    for cell in old_rows[idx].cells:
                        for p in cell.paragraphs:
                            for run in p.runs: run.font.color.rgb = COLOR_RED
            if len(new_rows) > min_rows:
                for idx in range(min_rows, len(new_rows)):
                    for cell in new_rows[idx].cells:
                        for p in cell.paragraphs:
                            for run in p.runs: run.font.color.rgb = COLOR_BLUE

def fit_table_to_width(tbl_xml, target_width_dxa):
    tblGrid = tbl_xml.find(qn('w:tblGrid'))
    col_widths = []
    if tblGrid is not None:
        for gridCol in tblGrid.findall(qn('w:gridCol')):
            w_val = gridCol.get(qn('w:w'))
            col_widths.append(float(w_val) if w_val else 1000.0)

    if not col_widths:
        first_row = tbl_xml.find(qn('w:tr'))
        if first_row is not None:
            num_cols = len(first_row.findall(qn('w:tc')))
            if num_cols > 0:
                col_widths = [1000.0] * num_cols

    if not col_widths:
        return

    total_orig_width = sum(col_widths)
    scale = target_width_dxa / total_orig_width
    scaled_widths = [int(w * scale) for w in col_widths]

    if tblGrid is not None:
        for gridCol in tblGrid.findall(qn('w:gridCol')):
            tblGrid.remove(gridCol)
        for sw in scaled_widths:
            new_col = OxmlElement('w:gridCol')
            new_col.set(qn('w:w'), str(sw))
            tblGrid.append(new_col)

    tblPr = tbl_xml.find(qn('w:tblPr'))
    if tblPr is not None:
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), str(target_width_dxa))

    for row in tbl_xml.findall(qn('w:tr')):
        col_idx = 0
        for tc in row.findall(qn('w:tc')):
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)

            span = 1
            gridSpan = tcPr.find(qn('w:gridSpan'))
            if gridSpan is not None:
                span_val = gridSpan.get(qn('w:val'))
                if span_val: span = int(span_val)

            if col_idx < len(scaled_widths):
                cell_width = sum(scaled_widths[col_idx : col_idx + span])
            else:
                cell_width = int((target_width_dxa / len(scaled_widths)) * span)

            col_idx += span
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:type'), 'dxa')
            tcW.set(qn('w:w'), str(cell_width))

def is_row_modified(tr_xml):
    for r in tr_xml.iter(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            color = rPr.find(qn('w:color'))
            if color is not None:
                val = color.get(qn('w:val'))
                if val and val.upper() in ('FF0000', '0000FF'):
                    return True
    return False

def replace_placeholder_in_element(container, placeholder, value):
    if hasattr(container, 'paragraphs'):
        for p in container.paragraphs:
            replace_placeholder_in_paragraph(p, placeholder, value)
    if hasattr(container, 'tables'):
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_placeholder_in_element(cell, placeholder, value)

def replace_placeholder_in_paragraph(p, placeholder, value):
    if placeholder not in p.text:
        return
    for run in p.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            return
    full_text = p.text.replace(placeholder, value)
    if p.runs:
        p.runs[0].text = full_text
        for run in p.runs[1:]:
            run.text = ""
    else:
        p.add_run(full_text)

def prune_paragraph_xml(p_xml):
    runs = p_xml.findall(qn('w:r'))
    if len(runs) <= 3:
        return
    modified_indices = set()
    for idx, r in enumerate(runs):
        rPr = r.find(qn('w:rPr'))
        if rPr is not None:
            color = rPr.find(qn('w:color'))
            if color is not None:
                val = color.get(qn('w:val'))
                if val in ('FF0000', '0000FF'):
                    modified_indices.add(idx)

    if not modified_indices:
        return

    keep_indices = set()
    for idx in modified_indices:
        keep_indices.add(idx)
        if idx > 0: keep_indices.add(idx - 1)
        if idx < len(runs) - 1: keep_indices.add(idx + 1)

    has_leading_ellipsis = 0 not in keep_indices
    has_trailing_ellipsis = (len(runs) - 1) not in keep_indices

    for idx, r in enumerate(runs):
        if idx not in keep_indices:
            p_xml.remove(r)

    if has_leading_ellipsis:
        first_kept_run = runs[min(keep_indices)]
        ellipsis_r = OxmlElement('w:r')
        ellipsis_t = OxmlElement('w:t')
        ellipsis_t.text = "... "
        ellipsis_r.append(ellipsis_t)
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '808080')
        rPr.append(color)
        rPr.append(OxmlElement('w:i'))
        ellipsis_r.append(rPr)
        first_kept_run.addprevious(ellipsis_r)

    if has_trailing_ellipsis:
        last_kept_run = runs[max(keep_indices)]
        ellipsis_r = OxmlElement('w:r')
        ellipsis_t = OxmlElement('w:t')
        ellipsis_t.text = " ..."
        ellipsis_r.append(ellipsis_t)
        rPr = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '808080')
        rPr.append(color)
        rPr.append(OxmlElement('w:i'))
        ellipsis_r.append(rPr)
        last_kept_run.addnext(ellipsis_r)

def prune_table_xml(tbl_xml):
    rows = tbl_xml.findall(qn('w:tr'))
    if len(rows) <= 3:
        return
    modified_indices = {idx for idx, tr in enumerate(rows) if idx > 0 and is_row_modified(tr)}
    if not modified_indices:
        return

    new_rows = [rows[0]]
    def create_ellipsis_row(template_tr):
        ell_tr = copy.deepcopy(template_tr)
        for cell in ell_tr.iter(qn('w:tc')):
            p = cell.find(qn('w:p'))
            if p is None:
                p = OxmlElement('w:p')
                cell.append(p)
            for r in list(p.findall(qn('w:r'))):
                p.remove(r)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '808080')
            rPr.append(color)
            rPr.append(OxmlElement('w:i'))
            r.append(rPr)
            t = OxmlElement('w:t')
            t.text = "..."
            r.append(t)
            p.append(r)
        return ell_tr

    last_kept_idx = 0
    for idx in sorted(modified_indices):
        if idx - last_kept_idx > 1:
            new_rows.append(create_ellipsis_row(rows[idx]))
        new_rows.append(rows[idx])
        last_kept_idx = idx

    if len(rows) - 1 - last_kept_idx > 0:
        new_rows.append(create_ellipsis_row(rows[-1]))

    for tr in rows: tbl_xml.remove(tr)
    for tr in new_rows: tbl_xml.append(tr)

def copy_group_blocks_to_cell(group, dest_cell, is_old):
    tc = dest_cell._tc
    for child in list(tc):
        if child.tag.endswith('Pr'):
            continue
        tc.remove(child)

    parts = [p.strip() for p in group['loc'].split(" > ")]
    for idx, part in enumerate(parts):
        p_hdr = OxmlElement('w:p')
        r_hdr = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rPr.append(OxmlElement('w:b'))

        sz_xml = OxmlElement('w:sz')
        sz_xml.set(qn('w:val'), '18')
        rPr.append(sz_xml)

        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        rFonts.set(qn('w:ascii'), '맑은 고딕')
        rFonts.set(qn('w:hAnsi'), '맑은 고딕')
        rPr.append(rFonts)

        r_hdr.append(rPr)
        t_hdr = OxmlElement('w:t')
        t_hdr.text = f"■ {part}" if idx == 0 else f"{'  ' * idx}└ {part}"
        r_hdr.append(t_hdr)
        p_hdr.append(r_hdr)
        tc.append(p_hdr)

    has_any_block = False
    for item in group['items']:
        src_block = item['old_block'] if is_old else item['new_block']
        if src_block is not None:
            has_any_block = True
            if isinstance(src_block, docx.text.paragraph.Paragraph):
                new_block_xml = copy.deepcopy(src_block._p)
                prune_paragraph_xml(new_block_xml)
                tc.append(new_block_xml)
            elif isinstance(src_block, docx.table.Table):
                new_block_xml = copy.deepcopy(src_block._tbl)
                prune_table_xml(new_block_xml)
                target_width_dxa = int(dest_cell.width.inches * 1440) - 300 if dest_cell.width else 4020
                if target_width_dxa <= 0: target_width_dxa = 4020
                fit_table_to_width(new_block_xml, target_width_dxa)
                tc.append(new_block_xml)

    if not has_any_block:
        p_none = OxmlElement('w:p')
        r_none = OxmlElement('w:r')
        rPr_none = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        rPr_none.append(rFonts)

        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '18')
        rPr_none.append(sz)

        color = OxmlElement('w:color')
        color.set(qn('w:val'), '808080')
        rPr_none.append(color)

        r_none.append(rPr_none)
        t_none = OxmlElement('w:t')
        t_none.text = "(없음)"
        r_none.append(t_none)
        p_none.append(r_none)
        tc.append(p_none)

    tc.append(OxmlElement('w:p'))

def format_cell_runs_font(cell, font_name='맑은 고딕', font_size_pt=9):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
    for table in cell.tables:
        for row in table.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size_pt)

def write_records_to_table(comparison_records, table, doc_subtype=""):
    grouped_records = []
    for record in comparison_records:
        if grouped_records and grouped_records[-1]['loc'] == record['loc']:
            grouped_records[-1]['items'].append(record)
        else:
            grouped_records.append({'loc': record['loc'], 'items': [record]})

    is_first_record = len(table.rows) == 2 and table.rows[1].cells[0].text == ""

    for r_idx, group in enumerate(grouped_records):
        row_cells = table.rows[1].cells if (is_first_record and r_idx == 0) else table.add_row().cells

        row_cells[0].text = doc_subtype
        set_cell_margins(row_cells[0])
        if row_cells[0].paragraphs[0].runs:
            run = row_cells[0].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            run.font.name = '맑은 고딕'
            run.font.bold = True

        copy_group_blocks_to_cell(group, row_cells[1], is_old=True)
        format_cell_runs_font(row_cells[1])
        set_cell_margins(row_cells[1])

        copy_group_blocks_to_cell(group, row_cells[2], is_old=False)
        format_cell_runs_font(row_cells[2])
        set_cell_margins(row_cells[2])

        row_cells[3].text = ""
        set_cell_margins(row_cells[3])

def find_preceding_table_title(blocks, current_idx):
    for idx in range(current_idx - 1, -1, -1):
        b = blocks[idx]
        if isinstance(b, docx.text.paragraph.Paragraph):
            text = b.text.strip()
            if not text: continue
            if any(re.match(p, text) for p in [r"^\s*(?:표|Table|<표|<Table)\s*\d+"]):
                return text
            if is_heading(b) and (re.match(r'^\d+\.\s', text) or re.match(r'^(?:■|◆|▶|◈)\s*', text)):
                break
    return None

def get_regulatory_location(blocks, current_idx):
    found_4_level, found_5_level, table_title = None, None, None
    b_curr = blocks[current_idx]
    if isinstance(b_curr, docx.table.Table):
        table_title = find_preceding_table_title(blocks, current_idx)

    for idx in range(current_idx - 1, -1, -1):
        b = blocks[idx]
        if not isinstance(b, docx.text.paragraph.Paragraph): continue
        text = b.text.strip()
        if not text: continue
        words = text.split()
        if not words: continue
        parts = words[0].rstrip('.').split('.')

        is_valid_reg = len(parts) >= 2 and all(re.match(r'^[a-zA-Z0-9\-]+$', p) for p in parts)
        if is_valid_reg:
            if len(parts) == 5 and not found_5_level: found_5_level = text
            elif len(parts) == 4 and not found_4_level: found_4_level = text
        if found_4_level and found_5_level: break

    path_nodes = []
    if found_4_level: path_nodes.append(found_4_level)
    if found_5_level: path_nodes.append(found_5_level)
    if table_title: path_nodes.append(table_title)

    if not path_nodes:
        for idx in range(current_idx - 1, -1, -1):
            b = blocks[idx]
            if isinstance(b, docx.text.paragraph.Paragraph):
                t = b.text.strip()
                if t and (re.match(r'^\d+\.', t) or re.match(r'^(?:■|◆|▶|◈)\s*', t)):
                    path_nodes.append(t)
                    break

    return " > ".join(path_nodes) if path_nodes else "문서 시작"

def remove_blue_color_from_document(doc):
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.color and run.font.color.rgb == RGBColor(0, 0, 255):
                run.font.color.rgb = RGBColor(0, 0, 0)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.color and run.font.color.rgb == RGBColor(0, 0, 255):
                            run.font.color.rgb = RGBColor(0, 0, 0)

def compare_and_modify_originals(doc1, doc2):
    remove_blue_color_from_document(doc1)

    blocks1 = list(iter_block_items(doc1))
    blocks2 = list(iter_block_items(doc2))

    blocks1_repr = [('p', b.text.strip()) if isinstance(b, docx.text.paragraph.Paragraph) else ('t', get_table_summary(b)) for b in blocks1]
    blocks2_repr = [('p', b.text.strip()) if isinstance(b, docx.text.paragraph.Paragraph) else ('t', get_table_summary(b)) for b in blocks2]

    matcher = difflib.SequenceMatcher(None, blocks1_repr, blocks2_repr)
    opcodes = matcher.get_opcodes()
    comparison_records = []

    for tag, i1, i2, j1, j2 in opcodes:
        if tag == 'equal': continue
        elif tag == 'insert':
            for j in range(j1, j2):
                comparison_records.append({'loc': get_regulatory_location(blocks2, j), 'old_block': None, 'new_block': blocks2[j], 'type': 'insert'})
        elif tag == 'delete':
            for i in range(i1, i2):
                comparison_records.append({'loc': get_regulatory_location(blocks1, i), 'old_block': blocks1[i], 'new_block': None, 'type': 'delete'})
        elif tag == 'replace':
            old_subset, new_subset = blocks1[i1:i2], blocks2[j1:j2]
            paired = match_replaced_ranges(old_subset, new_subset)
            for old_b, new_b in paired:
                loc = get_regulatory_location(blocks2, j1 + new_subset.index(new_b)) if new_b else get_regulatory_location(blocks1, i1 + old_subset.index(old_b))
                comparison_records.append({'loc': loc, 'old_block': old_b, 'new_block': new_b, 'type': 'replace'})

    for tag, i1, i2, j1, j2 in reversed(opcodes):
        if tag in ('equal', 'insert'): continue
        elif tag == 'delete':
            for i in range(i1, i2): color_entire_block(blocks1[i], COLOR_RED)
        elif tag == 'replace':
            if (i2 - i1 == 1) and (j2 - j1 == 1) and (blocks1_repr[i1][0] == blocks2_repr[j1][0]):
                b1, b2 = blocks1[i1], blocks2[j1]
                if isinstance(b1, docx.text.paragraph.Paragraph):
                    diff_runs = get_char_diff(b1.text.strip(), b2.text.strip())
                    apply_old_inline_diff(b1, diff_runs)
                else: compare_and_mark_tables(b1, b2)
            else:
                for i in range(i1, i2): color_entire_block(blocks1[i], COLOR_RED)

    for tag, i1, i2, j1, j2 in reversed(opcodes):
        if tag in ('equal', 'delete'): continue
        elif tag == 'insert':
            for j in range(j1, j2): color_entire_block(blocks2[j], COLOR_BLUE)
        elif tag == 'replace':
            if (i2 - i1 == 1) and (j2 - j1 == 1) and (blocks1_repr[i1][0] == blocks2_repr[j1][0]):
                b1, b2 = blocks1[i1], blocks2[j1]
                if isinstance(b2, docx.text.paragraph.Paragraph):
                    diff_runs = get_char_diff(b1.text.strip(), b2.text.strip())
                    apply_new_inline_diff(b2, diff_runs)
            else:
                for j in range(j1, j2): color_entire_block(blocks2[j], COLOR_BLUE)

    return comparison_records

@app.get("/ping")
def ping():
    return {"status": "ok", "message": "Server awake!"}

@app.post("/compare")
async def compare_documents(
    old_file: UploadFile = File(...),
    new_file: UploadFile = File(...),
    template_file: Optional[UploadFile] = File(None),
    product_name: Optional[str] = Form(""),
    doc_subtype: Optional[str] = Form(""),
    authorization: str = Header(None)
):
    if authorization != f"Bearer {MY_API_KEY}":
        raise HTTPException(status_code=401, detail="인증 실패")

    old_bytes = await old_file.read()
    new_bytes = await new_file.read()

    doc1 = docx.Document(io.BytesIO(old_bytes))
    doc2 = docx.Document(io.BytesIO(new_bytes))

    comparison_records = compare_and_modify_originals(doc1, doc2)

    if template_file:
        template_bytes = await template_file.read()
        doc_table = docx.Document(io.BytesIO(template_bytes))
        table = doc_table.tables[0]
    else:
        doc_table = docx.Document()
        title_p = doc_table.add_paragraph()
        title_run = title_p.add_run(f"변경대비표 - {product_name}" if product_name else "변경대비표")
        title_run.font.size = Pt(14)
        title_run.font.bold = True

        table = doc_table.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        set_table_borders(table)

        hdr_row = table.rows[0]
        headers = ["위치", "구버전", "신버전", "비고"]
        for idx, text in enumerate(headers):
            cell = hdr_row.cells[idx]
            cell.text = text
            set_cell_background(cell, HEX_HEADER_BG)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            if p.runs:
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(10)

    if product_name:
        clean_product_name = re.sub(r"\s+", " ", product_name).strip()
        replace_placeholder_in_element(doc_table, "%제품명%", clean_product_name)
        for section in doc_table.sections:
            if section.header is not None:
                replace_placeholder_in_element(section.header, "%제품명%", clean_product_name)

    if comparison_records:
        write_records_to_table(comparison_records, table, doc_subtype)
        col_widths = [Inches(1.5), Inches(3.0), Inches(3.0), Inches(1.0)]
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

    output_stream = io.BytesIO()
    doc_table.save(output_stream)
    output_stream.seek(0)

    return StreamingResponse(
        output_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=comparison_result.docx"}
    )
